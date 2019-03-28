import os
import logging
from datetime import datetime

from cryptography import x509
from cryptography.hazmat.backends import default_backend
from cryptography.x509 import ExtensionOID, NameOID
from docx import Document


class DERBasedGenerator:
    def __init__(self, template_path, certs_path, output_path):
        self.template_path = template_path
        self.certs_path = certs_path
        self.output_folder_path = output_path
        # surname, familyname, cnp
        self.cnp = [
            ("Ionela-Doriana", "Martin", "1111111111118"),
        ]

    def generate(self):
        certs = []

        if not os.path.exists(self.output_folder_path):
            os.mkdir(self.output_folder_path)

        for cert_file in os.listdir(self.certs_path):
            if cert_file.lower().endswith(".cer"):
                certs.append(os.path.join(self.certs_path, cert_file))

        print(
            "Vor fi tiparite in syntaxa Python, informatiile care nu s-au gasit in dictionarul intern."
        )

        for cpath in certs:
            with open(cpath, "rb") as cert_file:
                der = cert_file.read()

                cert = x509.load_der_x509_certificate(der, default_backend())

                email = cert.extensions.get_extension_for_oid(
                    ExtensionOID.SUBJECT_ALTERNATIVE_NAME
                ).value.get_values_for_type(x509.RFC822Name)[0]

                not_before = cert.not_valid_before.strftime("%c")

                not_after = cert.not_valid_after.strftime("%c")

                sn = hex(cert.serial_number).strip("0x")

                serial_number = ":".join(
                    a + b for a, b in zip(sn[::2], sn[1::2])
                ).upper()

                issuer = cert.issuer.get_attributes_for_oid(x509.OID_COMMON_NAME)[
                    0
                ].value

                # subject_name = cert.subject.get_attributes_for_oid(
                #     x509.OID_COMMON_NAME
                # )[0].value

                subject_surname = cert.subject.get_attributes_for_oid(
                    x509.OID_GIVEN_NAME
                )[0].value

                subject_familyname = cert.subject.get_attributes_for_oid(
                    x509.OID_SURNAME
                )[0].value

                subject_cnp = ""

                out_filename = (
                    self.output_folder_path
                    + "/"
                    + subject_familyname
                    + " "
                    + subject_surname
                    + " "
                    + sn
                    + ".docx"
                )

                gasit = False

                for c in self.cnp:
                    if c[0] == subject_surname and c[1] == subject_familyname:
                        subject_cnp = c[2]
                        gasit = True

                if not gasit:
                    print(
                        '("'
                        + subject_surname
                        + '", "'
                        + subject_familyname
                        + '", "cnp_aici"),'
                    )

                # print(subject_familyname + " " + subject_surname + ", " + not_before + " - " + not_after + ", " + serial_number + ", " + issuer + " -> " + out_filename)

                document = Document(self.template_path)

                try:
                    # nume:
                    document.tables[1].cell(0, 1).text = subject_surname
                    # prenume:
                    document.tables[1].cell(1, 1).text = subject_familyname
                    # email:
                    document.tables[1].cell(2, 1).text = email
                    # serial:
                    document.tables[1].cell(3, 1).text = serial_number
                    # autoritate:
                    document.tables[1].cell(4, 1).text = issuer
                    # cnp:
                    document.tables[1].cell(5, 1).text = subject_cnp
                    # valabilitate:
                    document.tables[1].cell(6, 1).text = not_before + " - " + not_after

                    # salvare:
                    document.save(out_filename)
                except Exception as e:
                    logging.error(f"Failed to edit template: {e}")


if __name__ == "__main__":
    g = DERBasedGenerator("./template.docx", "./input", "./output")
    g.generate()
